from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
MD_OUT = OUT_DIR / "final_internship_report.md"
DOCX_OUT = OUT_DIR / "final_internship_report.docx"


TITLE = "车牌号识别项目最终实习报告"
SUBTITLE = "基于 YOLO 车牌检测与 CRNN + CTC 整牌识别的工程实践"
META = [
    ("实习项目", "中国车牌号检测与识别系统"),
    ("实习时间", "2026 年 7 月"),
    ("实习人", "XXX"),
    ("指导老师", "XXX"),
    ("项目目录", "license-plate-recognition-week3"),
]


SECTIONS = [
    (
        "一、实习项目概述",
        [
            "本次实习围绕中国车牌号自动识别任务展开，目标是构建一套能够从车辆图片中定位车牌并输出完整车牌号码的识别系统。项目早期主要解决“裁剪后的车牌图片到完整车牌号”的识别问题，后续进一步扩展为“整车图片到车牌检测、裁剪、识别和可视化输出”的完整流程。",
            "项目使用 CCPD 数据集作为主要训练与评估数据来源。CCPD 文件名中包含车牌位置、四点标注和字符编码信息，因此能够同时支持车牌检测数据集构建、车牌裁剪图识别训练以及后续错误分析。经过多轮实验，当前系统推荐使用 YOLO 负责车牌定位，使用 CRNN + CTC 负责完整车牌字符串识别。",
            "在整个实习过程中，我主要承担数据整理、训练数据构建、模型训练与微调、推理流程串联、结果评估和错误分析等工作。最终项目从单一的裁剪车牌识别能力，升级为具备整车图片输入、车牌检测、车牌裁剪、整牌识别和可视化输出能力的完整工程。",
        ],
    ),
    (
        "二、实习工作内容",
        [
            "1. 数据整理与标签解析。负责整理 CCPD 数据流程，解析原始图片文件名中的车牌号码、bbox、四点坐标和数据划分信息，生成统一的 labels.csv 文件。最终整理出 351,977 条样本记录，其中训练集 246,370 条、验证集 52,827 条、测试集 52,780 条。",
            "2. YOLO 检测数据集构建。基于 labels.csv 中的 bbox 字段，将 CCPD 标注转换为 YOLO 所需的 class_id、x_center、y_center、width、height 格式，并按照 train、val 划分组织 images 和 labels 目录。为节省磁盘空间，数据集组织默认采用软链接方式，避免重复复制大规模图片。",
            "3. YOLO 车牌检测器训练与评估。使用 YOLOv8n 作为基础模型，完成 5k 和 20k 规模实验，训练并对比 plate_detector_5k、plate_detector_20k-3 等模型，最终选用 plate_detector_20k-3 作为主要检测模型，并将最佳权重整理到 models/plate_detector.pt。",
            "4. CRNN + CTC 整牌识别模型微调。围绕裁剪车牌图到完整车牌号的序列识别任务，对 CRNN 模型进行长周期微调，将训练轮数由 20 epoch 扩展到 80 epoch，batch size 从 32 提升到 64，并引入 StepLR(step=15, gamma=0.5) 学习率衰减策略。",
            "5. 端到端推理流程开发。新增和完善 detect_plate_yolo.py、detect_and_recognize.py 等脚本，打通整车图片输入、车牌定位、车牌裁剪、CRNN 识别、结果绘制和文件保存流程，使项目具备从原始车辆图片直接输出车牌号的能力。",
            "6. YOLO 字符检测对照实验。实现 char_detector.pt 的训练与评估流程，尝试通过字符框检测、横向排序和车牌规则后处理拼接车牌号，并整理 yolo_chars_v2 的检测指标、端到端准确率和错误样本，为后续优化提供依据。",
            "7. 结果归档与错误分析。整理训练曲线、预测明细、检测可视化图、裁剪车牌图、整车识别结果图和错误样本文件，形成当前模型总结、周工作报告、YOLO 字符优化方案和最终实习报告等文档。",
        ],
    ),
    (
        "三、个人工作量与交付成果",
        [
            "本次实习不是单一模型调用或简单脚本运行，而是围绕一个完整视觉识别项目完成了从数据到模型、从训练到部署推理、从指标统计到错误分析的闭环工作。个人工作量主要体现在数据规模处理、模型多轮训练、工程脚本开发、评估结果沉淀和最终文档整理几个方面。",
            "在数据侧，我完成了 CCPD 标签解析和 YOLO 训练数据转换，使原始数据能够被检测模型和识别模型稳定使用；在模型侧，我分别推进 YOLO 车牌检测、CRNN 整牌识别和 YOLO 字符检测三条实验线；在工程侧，我将独立的检测与识别模块串联为可运行的端到端流程，并保留结果图和预测明细，便于复现和汇报。",
        ],
    ),
    (
        "四、技术路线与系统结构",
        [
            "系统主流程采用两阶段结构。第一阶段由 YOLO 车牌检测器在整车图片中定位车牌区域；第二阶段将检测到的车牌区域裁剪后送入 CRNN + CTC 识别器，直接输出完整车牌字符串。该路线兼顾了目标检测模型对复杂背景的定位能力，以及序列识别模型对整牌字符顺序的建模能力。",
            "当前推荐主流程如下：整车图片 -> YOLO 车牌检测 -> 裁剪车牌区域 -> CRNN + CTC 整牌识别 -> 输出车牌号码和可视化结果。",
            "项目中也保留了纯 YOLO 字符检测路线，即整车图片 -> YOLO 车牌检测 -> 裁剪车牌区域 -> YOLO 字符检测 -> 按 x_center 排序 -> 规则后处理 -> 拼接车牌号。该方案可解释性较强，但受伪字符框标签、多检漏检和序列规则限制，整牌准确率低于 CRNN 主流程。",
        ],
    ),
    (
        "五、数据集与模型训练",
        [
            "项目主要数据来自 CCPD，共整理出 351,977 张样本，其中训练集 246,370 张，验证集 52,827 张，测试集 52,780 张。训练过程中严格区分 train、val、test，车牌检测训练只使用 train 和 val，避免测试集泄漏。",
            "YOLO 车牌检测器使用 YOLOv8n 作为基础模型，目标类别为 license_plate。主要训练目录为 runs/detect/plate_detector_20k-3，输入尺寸为 416，训练 50 epoch，部署权重为 models/plate_detector.pt。",
            "CRNN + CTC 识别器由 CNN 特征提取层、双向 LSTM 序列建模层、线性分类层和 CTC 解码组成。第八周微调将 batch size 从 32 提升到 64，将训练轮数从 20 扩展到 80，并采用阶段式学习率衰减。微调后识别模块成为当前系统中最稳定的核心组件。",
        ],
    ),
    (
        "六、实验结果与模型成果",
        [
            "从实验结果看，车牌检测器已经能够稳定找到车牌位置。plate_detector_20k-3 在验证集上 mAP50 达到 99.46% 左右，precision 和 recall 均接近 99.9%。mAP50-95 为 79.21% 至 79.34%，说明在严格 IoU 阈值下，检测框精细定位仍有继续提升空间。",
            "CRNN + CTC 识别器在第八周微调后取得当前最佳结果：字符准确率达到 99.7%，整牌准确率达到 98% 以上。相比此前 CCPD test split 上字符准确率 98.94%、整牌准确率 94.43% 的基线，整牌准确率至少提升 3.57 个百分点，错误率从约 5.57% 降至 2% 以下，说明本轮微调明显提升了完整车牌输出的稳定性。",
            "YOLO 字符检测方案在 yolo_chars_v2 实验中 mAP50 达到 76.668%，但端到端整牌准确率约为 76.81%，明显低于 CRNN 主流程。主要原因在于字符框标签来自宽度均分伪标注，不是真实字符边界，容易造成多检、漏检或排序错误。",
            "综合模型成果看，YOLO 车牌检测 + CRNN 整牌识别是当前最可靠的主流程。检测模型解决了整车图中车牌位置未知的问题，识别模型解决了完整车牌字符串输出的问题，两者串联后使系统从“只能识别裁剪车牌图”提升为“可以处理整车图片”的可用工程。",
        ],
    ),
    (
        "七、工程实现成果",
        [
            "本次实习最终形成了可运行、可训练、可评估的车牌识别工程。项目中包含数据准备、模型训练、模型评估、单图预测、整车图检测识别、YOLO 字符识别对照实验和错误分析等模块。",
            "主要代码包括 prepare_ccpd.py、prepare_yolo_plate_dataset.py、detect_plate_yolo.py、detect_and_recognize.py、train.py、evaluate.py、predict.py、prepare_yolo_char_dataset.py、train_yolo_char.py、recognize_plate_yolo.py、evaluate_yolo_char.py、plate_rules.py 和 plate_geometry.py 等。",
            "主要模型文件包括 models/plate_detector.pt、models/license_plate_crnn_best.pth 和 models/char_detector.pt。主要输出结果包括训练曲线、检测可视化图、裁剪车牌图、整车识别结果图、预测明细表和错误样本分析文件。",
        ],
    ),
    (
        "八、问题分析与改进方向",
        [
            "1. 真实场景端到端评估仍需加强。当前 CCPD 指标较好，但自摄图片存在光照、角度、模糊、遮挡和压缩质量差异，需要建立独立真实测试集，人工标注车牌号并统计端到端准确率。",
            "2. 当前裁剪主要依赖矩形 bbox。对于倾斜或透视变形明显的车牌，矩形裁剪可能导致字符形态变形，后续应优先利用 CCPD 四点标注或关键点估计进行透视矫正。",
            "3. 多车牌场景支持不足。当前系统默认选择置信度最高的一个车牌框，尚未系统支持一张图中多个车牌逐个识别。",
            "4. 特殊车牌类型覆盖有限。当前字符表主要覆盖常规省份简称、字母和数字，对新能源车牌、警牌、学牌、挂车、港澳车牌等复杂场景支持仍需扩展。",
            "5. YOLO 字符检测标签质量有限。若继续推进纯 YOLO 字符识别方案，应优先改进伪字符框生成方式，或引入真实字符级标注，提高字符定位可靠性。",
        ],
    ),
    (
        "九、实习收获与总结",
        [
            "通过本次实习，我完整参与并实现了一个从数据整理、模型训练、实验评估到工程推理的计算机视觉项目。项目不仅涉及深度学习模型训练，还包含大规模数据格式转换、训练集构建、模型权重归档、推理脚本封装、结果可视化和错误分析等工程环节。",
            "在技术层面，我加深了对目标检测、序列识别和 OCR 任务的理解。YOLO 适合解决整车图中的车牌定位问题，CRNN + CTC 适合解决车牌字符序列识别问题，两者组合能够形成较稳定的端到端系统。同时，YOLO 字符检测实验也让我认识到，模型结构之外，标注质量和后处理规则对最终效果同样重要。",
            "在工程层面，我体会到一个可复现项目需要清晰的数据入口、稳定的训练脚本、明确的评估指标和可追踪的结果文件。最终，本项目已经具备完整的车牌识别主流程，在 CCPD 标准场景下达到较高识别效果，也为后续真实场景优化、多车牌识别、透视矫正和特殊车牌扩展打下了基础。",
        ],
    ),
]


TABLES = {
    "个人工作量概览": [
        ["工作模块", "完成内容", "量化成果"],
        ["数据整理", "解析 CCPD 文件名，生成统一 labels.csv，保留车牌号、bbox、points 和 split 信息", "351,977 条样本记录"],
        ["检测数据构建", "将 CCPD bbox 转换为 YOLO 格式，组织 train/val 图片与标签目录", "246,370 train / 52,827 val"],
        ["检测模型训练", "训练并对比 YOLOv8n 车牌检测模型，归档最佳权重", "mAP50 约 99.46%"],
        ["识别模型微调", "将 CRNN 训练扩展到 80 epoch，batch size 提升至 64，引入 StepLR 衰减", "整牌准确率 98%+，字符准确率 99.7%"],
        ["端到端流程", "开发检测、裁剪、识别、绘图脚本", "支持整车图直接输出车牌号"],
        ["分析与文档", "整理训练曲线、预测明细、错误样本、周报和最终报告", "形成可复现实验材料"],
    ],
    "数据集划分": [
        ["split", "样本数"],
        ["train", "246,370"],
        ["val", "52,827"],
        ["test", "52,780"],
        ["合计", "351,977"],
    ],
    "核心模型": [
        ["模型", "权重文件", "作用", "当前定位"],
        ["YOLO 车牌检测器", "models/plate_detector.pt", "定位整车图中的车牌框", "主流程检测模块"],
        ["CRNN + CTC 识别器", "models/license_plate_crnn_best.pth", "裁剪车牌图到完整车牌号", "主流程识别模块"],
        ["YOLO 字符检测器", "models/char_detector.pt", "检测单字符并排序拼接", "对照实验和辅助分析"],
    ],
    "关键性能指标": [
        ["模块", "指标", "结果"],
        ["YOLO 车牌检测", "mAP50", "99.46%"],
        ["YOLO 车牌检测", "mAP50-95", "79.21% 至 79.34%"],
        ["CRNN + CTC", "字符准确率", "99.7%"],
        ["CRNN + CTC", "整牌准确率", "98%+"],
        ["YOLO 字符检测", "端到端整牌准确率", "76.81%"],
    ],
    "模型成果对比": [
        ["模型/方案", "优化前或对照结果", "优化后结果", "结论"],
        ["CRNN + CTC 整牌识别", "整牌准确率 94.43%，字符准确率 98.94%", "整牌准确率 98%+，字符准确率 99.7%", "主识别模型显著提升"],
        ["YOLO 车牌检测", "5k 实验 mAP50 99.49%，mAP50-95 74.95%", "20k-3 实验 mAP50 99.46%，mAP50-95 79.34%", "严格 IoU 下定位更稳"],
        ["YOLO 字符检测", "作为对照实验路线", "端到端整牌准确率 76.81%", "保留为辅助分析，不替代 CRNN"],
    ],
}


FIGURES = [
    ("整体性能摘要", ROOT / "results/work_report_figures/pipeline_performance_summary.png"),
    ("CRNN 训练曲线", ROOT / "results/work_report_figures/crnn_training_summary.png"),
    ("YOLO 训练曲线", ROOT / "results/work_report_figures/yolo_training_summary.png"),
]


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    run.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_font(run, 11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 13, bold=True, color="2E74B5")
    return p


def add_docx_table(doc, title, rows):
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value, bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(cell, "F2F4F7")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_markdown():
    lines = [f"# {TITLE}", "", SUBTITLE, ""]
    for key, value in META:
        lines.append(f"- **{key}：**{value}")
    lines.append("")
    for heading, paragraphs in SECTIONS:
        lines.extend([f"## {heading}", ""])
        for para in paragraphs:
            lines.extend([para, ""])
        if heading == "三、个人工作量与交付成果":
            rows = TABLES["个人工作量概览"]
            lines.extend(["### 个人工作量概览", ""])
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("|" + "|".join(["---"] * len(rows[0])) + "|")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
        if heading == "五、数据集与模型训练":
            for title in ("数据集划分", "核心模型"):
                rows = TABLES[title]
                lines.extend([f"### {title}", ""])
                lines.append("| " + " | ".join(rows[0]) + " |")
                lines.append("|" + "|".join(["---"] * len(rows[0])) + "|")
                for row in rows[1:]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
        if heading == "六、实验结果与模型成果":
            for title in ("关键性能指标", "模型成果对比"):
                rows = TABLES[title]
                lines.extend([f"### {title}", ""])
                lines.append("| " + " | ".join(rows[0]) + " |")
                lines.append("|" + "|".join(["---"] * len(rows[0])) + "|")
                for row in rows[1:]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    header_p = section.header.paragraphs[0]
    header_p.text = "车牌号识别项目最终实习报告"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.runs[0].font.size = Pt(9)
    header_p.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    section.footer.paragraphs[0].text = ""

    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run(TITLE)
    set_font(title_run, 24, bold=True, color="0B2545")

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(20)
    subtitle_run = subtitle_p.add_run(SUBTITLE)
    set_font(subtitle_run, 13, color="555555")

    meta_table = doc.add_table(rows=len(META), cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    set_table_borders(meta_table)
    for idx, (key, value) in enumerate(META):
        set_cell_text(meta_table.cell(idx, 0), key, bold=True)
        set_cell_text(meta_table.cell(idx, 1), value)
        shade_cell(meta_table.cell(idx, 0), "F2F4F7")
    doc.add_page_break()

    for heading, paragraphs in SECTIONS:
        add_heading(doc, heading, level=1)
        for para in paragraphs:
            add_paragraph(doc, para)
        if heading == "三、个人工作量与交付成果":
            add_docx_table(doc, "个人工作量概览", TABLES["个人工作量概览"])
        if heading == "五、数据集与模型训练":
            add_docx_table(doc, "数据集划分", TABLES["数据集划分"])
            add_docx_table(doc, "核心模型", TABLES["核心模型"])
        if heading == "六、实验结果与模型成果":
            add_docx_table(doc, "关键性能指标", TABLES["关键性能指标"])
            doc.add_page_break()
            add_docx_table(doc, "模型成果对比", TABLES["模型成果对比"])
            for caption, path in FIGURES:
                if path.exists():
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.paragraph_format.space_before = Pt(6)
                    cap.paragraph_format.space_after = Pt(3)
                    cap_run = cap.add_run(caption)
                    set_font(cap_run, 9.5, color="555555")
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(8)
                    run = p.add_run()
                    run.add_picture(str(path), width=Inches(5.35))

    doc.save(DOCX_OUT)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_markdown()
    build_docx()
    print(MD_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
